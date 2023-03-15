def textget():
    import docx
    import olefile
    from docx import Document
    from docx.shared import Inches
    #doc = docx.Document("lesson.doc")
    import aspose.words as aw
    from pdf2docx import Converter


    fileNames = [ "lesson.doc"]
    output = aw.Document()
    # Remove all content from the destination document before appending.
    output.remove_all_children()

    for fileName in fileNames:
        input = aw.Document(fileName)
        # Append the source document to the end of the destination document.
        output.append_document(input, aw.ImportFormatMode.KEEP_SOURCE_FORMATTING)

    output.save("lesson.docx")

    #doc1 = aw.Document("lesson.doc")
    #doc1.save("lesson.docx")
    doc = docx.Document("lesson.docx")
    columnsName=''
    Table={}       
    WeekDays=['monday', 'tuesday', 'wednesday', 'thursday','friday']  
    i=-1      
    for table in doc.tables:
        for row in table.columns:
            for cell in row.cells:
                if(cell.text=='8-А'):
                    columnsName='8-А'
                    i=0
                if(cell.text=='8-Б'):
                    columnsName='8-Б'
                    i=0
                if(cell.text=='9-А'):
                    columnsName='9-А'
                    i=0
                if(cell.text=='9-Б'):
                    columnsName='9-Б'
                    i=0
                if(cell.text=='10-А'):
                    columnsName='10-А'
                    i=0
                if(cell.text=='10-Б'):
                    columnsName='10-Б'
                    i=0
        
                if(cell.text=='10-В'):
                    columnsName='10-В'
                    i=0
                if(cell.text=='11-А'):
                    columnsName='11-А'
                    i=0
                if(cell.text=='11-Б'):
                    columnsName='11-Б'
                    i=0
                if(cell.text=='11-В'):
                    columnsName='11-В'
                    i=0
                if(i==-1):
                        continue
                classList=Table.get(columnsName)
                nDay=i//10
                if(classList==None):
                        Table.update({columnsName: {}})        
                else:
                        currentDay=WeekDays[nDay]
                        dayList=classList.get(currentDay)
                        textTable=cell.text.replace("\n", "")
                        if(dayList==None):
                                classList.update({currentDay: [textTable]})
                        else:
                                dayList.append(cell.text)
                                classList.update({currentDay: dayList})
                        Table.update({columnsName: classList})
                        i=i+1
                
     
    for key, classList in Table.items():
        for keyDay, dayList in classList.items():
            for i in range(len(dayList)):
                if(dayList[i]==''):
                        dayList[i]="Вiкно"
                        
                changeN=dayList[i]
                changeN=changeN.replace("\n","")
                dayList[i]=changeN

    return(Table)
print(textget())